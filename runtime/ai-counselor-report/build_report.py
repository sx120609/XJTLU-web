from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = BASE_DIR / "ai-counselor-report.docx"

IMG_HOMEPAGE = BASE_DIR / "01-homepage-workbench-concept.png"
IMG_ARCH = BASE_DIR / "02-capability-architecture.png"
IMG_FORUM = BASE_DIR / "03-forum-boundary.png"


TITLE = "关于我们 AI 辅导员产品形态与能力边界的建议"
SUBTITLE = "围绕产品结构、能力组织和论坛边界的阶段性整理"


BODY_FONT = "Microsoft YaHei"
HEADING_COLOR = RGBColor(32, 80, 128)
TEXT_COLOR = RGBColor(32, 35, 42)
MUTED_COLOR = RGBColor(100, 108, 122)


SECTIONS = [
    (
        "一、我们的基本判断",
        [
            "如果我们要把 AI 辅导员做成一个真正能长期使用的系统，那它的核心就不应该只是“会聊天”，而应该是“能帮学生把事情往下办”。学生进入这个页面，很多时候不是为了和 AI 对话本身，而是为了更快完成具体任务，比如查课表、看通知、找流程、交材料、联系部门。",
            "所以从产品定位上看，我们更适合把它理解成一个学生服务工作台，而不是一个放大版聊天页。",
        ],
    ),
    (
        "二、现有页面的主要问题",
        [
            "现在这个版本最大的问题，不是视觉风格本身，而是结构重心有点偏。",
        ],
        [
            "中间聊天区太重，占了最核心的位置，但没有承接最常见的校园任务。",
            "左侧历史会话、右侧快捷入口、常见问题，本质上都在做导航，信息比较散，也有重复。",
            "页面缺少“和我有关”的内容。学生点进来之后，最先关心的通常是今天有没有课、最近有没有考试、有没有材料没交、最近有没有新通知，而不是一句欢迎语。",
            "AI 回答完之后，下一步动作不够清楚。告诉你流程是一回事，能不能顺着点进去把事办完是另一回事。",
            "页面里的可信度表达还不够强。对学校场景来说，来源、更新时间、适用对象这些信息其实很重要。",
        ],
        [
            "从这个角度看，现有版本更像是一个 AI 展示页，还不像一个真正能承接学生服务的 AI 辅导员。",
        ],
    ),
    (
        "三、我们建议的产品定位",
        [
            "我们更建议把 AI 辅导员定位成“学生服务工作台 + 对话式引导层”。",
            "这里的重点是，首页优先承接任务，对话能力负责做补充。AI 最有价值的地方，不是单纯生成一段回答，而是：",
        ],
        [
            "帮学生判断问题属于哪类事务。",
            "帮学生把分散信息整理成更好理解的结果。",
            "帮学生进入正确的办理入口。",
            "在复杂、模糊、跨系统的问题上继续追问和引导。",
        ],
        [
            "如果按这个方向做，AI 辅导员会更像一个可持续建设的校内系统，而不是一个阶段性的功能展示。",
        ],
    ),
    (
        "四、首页重构方向",
        [
            "这次重画首页的思路，就是把页面从“聊天中心”改成“任务中心”。",
            "首页更适合优先展示几类学生最关心的信息，比如“今日课表”“最近考试”“待交材料”“最新公告”。AI 输入框仍然保留，但不再独占页面中心，而是作为“描述复杂问题”或“补充说明”的入口。再往下，直接放一组高频任务入口，比如“查课表”“查成绩”“办事流程”“提交材料”“联系部门”“校园公告”“宿舍电费”“常见问题”。",
            "右侧区域也不建议继续放静态介绍，而更适合放“我的待办”“办事进度”“可信来源”“最近咨询”这类真正有用的动态信息。",
            "这样改的好处是，学生一进来就能明白两件事：这个系统知道我现在最可能需要什么；它不只是会回答问题，而是真的能帮我推进事情。",
        ],
    ),
    (
        "五、我们真正值得吸收的部分",
        [
            "如果我们要吸收现有项目中的一些能力，我觉得重点不应该放在照搬页面，而应该放在吸收背后的组织方式。",
            "更值得吸收的部分包括：",
        ],
        [
            "把分散校园服务整理成统一入口的思路。",
            "把高频任务拆成明确功能模块的思路。",
            "教务、公告、问卷、文件收集、云盘这类轻能力的组织方式。",
            "哪些内容可公开、哪些必须登录、哪些只能本人查看的权限边界。",
            "AI 回答之后，如何把用户送到真正的办理入口，而不是停留在一段文字里。",
        ],
        [
            "简单说，真正有价值的不是一个站点皮肤，而是“服务整合能力”和“任务落地能力”。",
        ],
    ),
    (
        "六、关于论坛能力，能不能吸收",
        [
            "我觉得可以吸收一部分，但不建议把完整论坛直接并进 AI 辅导员。",
            "论坛类内容的价值很明显。它能补足纯官方知识库覆盖不到的实际经验，比如某个流程里容易卡住的地方、材料提交时常见的问题、学生真实会怎么提问。它也能持续沉淀新的问题场景，让系统不是只会回答预设 FAQ，而是能不断接近真实使用。",
            "但问题也同样明显。一旦进入官方系统，论坛就不再只是社区功能，而会带上很强的治理属性。开放灌水、强匿名吐槽、未经整理的自由讨论、二手交易这类内容，如果直接并进来，风险会很高，也会把产品重心带偏。",
            "所以更合适的方式不是“把完整论坛搬进来”，而是把其中最有价值、最可治理的部分提炼出来，做成一个经验问答层或案例知识层。",
        ],
        [
            "已解决问题。",
            "办事经验整理。",
            "高频咨询归档。",
            "常见踩坑提醒。",
            "经审核后的学生经验补充。",
        ],
        [
            "这里最关键的一点，是要把“官方信息”和“学生经验”明确分层。官方信息优先，学生经验单独标注，来源要清晰，AI 如果引用经验内容，也要能追溯到具体出处。这样论坛能力才能真正成为补充，而不是干扰。",
        ],
    ),
    (
        "七、我们的推进建议",
        [
            "从落地顺序上看，我会建议先把基础层做稳，再逐步扩展。",
        ],
        [
            "第一阶段先做好只读型能力，比如公告问答、流程解释、课表查询、成绩查询、考试信息、联系方式查询。",
            "第二阶段再接入任务型能力，比如材料提交、问卷填写、办事进度查看、常见业务入口跳转。",
            "第三阶段逐步加入个性化看板，根据学生身份、时间节点和任务状态展示待办事项。",
            "第四阶段再考虑接入经过治理的经验问答层，把论坛里真正有价值的部分沉淀进来。",
        ],
        [
            "这样推进，风险更可控，产品结构也更容易立住。",
        ],
    ),
    (
        "八、结论",
        [
            "我们的核心问题，其实不是首页要不要更好看，而是整个 AI 辅导员到底按什么思路来做。",
            "如果继续沿着“聊天页”去堆功能，它很容易停留在“会回答，但不太会办事”的状态。更合适的方向，是把它做成一个以任务、待办和可信信息为中心的学生服务工作台，再让 AI 去承担解释、引导和分发的工作。",
            "如果我们要吸收现有项目中的价值，重点应该放在服务整合、能力编排、权限边界和任务落地这些方面。至于论坛能力，可以吸收，但必须收边界、做治理，更适合沉淀成经验问答层，而不是直接引入一个完整社区。",
        ],
    ),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_sep)
    paragraph._p.append(fld_end)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("第 ")
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_COLOR
    add_page_number(p)
    run = p.add_run(" 页")
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_COLOR


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = HEADING_COLOR
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.1

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15


def add_title_block(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(TITLE)
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR

    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    run = sub.add_run(SUBTITLE)
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED_COLOR

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    note.paragraph_format.left_indent = Cm(0.2)
    shade_paragraph(note, "EEF4FB")
    run = note.add_run("摘要：这份材料主要围绕 AI 辅导员的产品定位、首页结构、能力组织以及论坛能力的吸收边界，提出一套更适合校内正式场景的建设思路。")
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_COLOR


def add_section(document: Document, title: str, paragraphs: list[str], bullets: list[str] | None = None, tail: list[str] | None = None) -> None:
    document.add_paragraph(title, style="Heading 1")
    for text in paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        set_east_asia_font(run, BODY_FONT)
    if bullets:
        for item in bullets:
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(item)
            set_east_asia_font(run, BODY_FONT)
    if tail:
        for text in tail:
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(text)
            set_east_asia_font(run, BODY_FONT)


def add_figure_section(document: Document, heading: str, image_path: Path, caption: str, lead: str) -> None:
    document.add_page_break()
    document.add_paragraph(heading, style="Heading 2")

    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(lead)
    set_east_asia_font(run, BODY_FONT)

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_before = Pt(8)
    picture.paragraph_format.space_after = Pt(6)
    picture.add_run().add_picture(str(image_path), width=Inches(6.15))

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    run = cap.add_run(caption)
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED_COLOR


def add_appendix_note(document: Document) -> None:
    document.add_page_break()
    document.add_paragraph("附：三张配图的使用建议", style="Heading 1")
    for text in [
        "首页概念图适合放在汇报前半部分，用来直观说明“为什么首页应该从聊天中心转成任务中心”。",
        "能力结构图适合放在中段，用来说明 AI 辅导员不是单一模型页面，而是展示层、AI 编排层和能力数据层共同组成的系统。",
        "论坛边界图适合放在后半部分，用来说明论坛能力可以吸收，但必须经过治理和重构，不能直接以完整社区形态并入。",
    ]:
        p = document.add_paragraph(style="List Number")
        run = p.add_run(text)
        set_east_asia_font(run, BODY_FONT)


def build_document() -> None:
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_title_block(document)

    for section in SECTIONS[:3]:
        add_section(document, *section)

    add_section(document, *SECTIONS[3])

    add_figure_section(
        document,
        "配图一：工作台式首页概念图",
        IMG_HOMEPAGE,
        "图 1  首页概念图：从“聊天中心”转向“任务中心”的工作台式首页",
        "这版首页不是把 AI 去掉，而是把 AI 放回到更合适的位置。页面先承接学生最关心的事情，再让对话能力去补足复杂问题，整体会更像一个真实可用的校内服务入口。",
    )

    add_section(document, *SECTIONS[4])

    add_figure_section(
        document,
        "配图二：建议能力结构",
        IMG_ARCH,
        "图 2  能力结构图：展示层、AI 编排层与能力数据层的分层关系",
        "如果我们后续要把更多能力接进来，最稳的方式不是不停给聊天框加功能，而是把页面展示、意图识别和底层能力调用拆开，这样系统扩展起来会更清楚。",
    )

    add_section(document, *SECTIONS[5])

    add_figure_section(
        document,
        "配图三：论坛能力的吸收边界",
        IMG_FORUM,
        "图 3  论坛边界图：保留经验价值，但避免把完整社区直接带入官方场景",
        "论坛最值得保留的是经验沉淀和真实问题，而不是所有原始讨论。把它转成“经验问答层”或“案例知识层”，既能补足官方知识库，又能把风险控制在更合理的范围内。",
    )

    add_section(document, *SECTIONS[6])
    add_section(document, *SECTIONS[7])
    add_appendix_note(document)
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
